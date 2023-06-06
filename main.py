import MTK2
from docx import Document
from docx.shared import RGBColor

document = Document('variant07.docx')
paragraphs = document.paragraphs

def main():
    # phrase = str(input("Enter quote to concealment: "))
    phrase = 'Крепкую дружбу и топором не разрубишь.'

    PhraseToMTK2 = MTK2.Encode(phrase)
    print('Text in MTK2:', PhraseToMTK2)
    print(len(PhraseToMTK2))
    PhraseToMTK2Dec = MTK2.Decode(PhraseToMTK2)
    print('Text in MTK2 Decode:', PhraseToMTK2Dec)      #проверка

    ArrLenString = []
    OpenText = ""

    print('len doc parag:', len(document.paragraphs))
    for paragraph in document.paragraphs:
        String = ""
        for run in paragraph.runs:
            for char in run.text:
                String += char
        OpenText += String
        ArrLenString.append(len(String))
    print('Open text:', OpenText)
    print('Len paragraphs:', ArrLenString)

    document.paragraphs.clear()

    countOfTextEl = 0
    for id_paragraph in range(len(document.paragraphs)):
        document.paragraphs[id_paragraph].clear()
        for id_runs in range(ArrLenString[id_paragraph]):
            run = document.paragraphs[id_paragraph].add_run(OpenText[countOfTextEl]) #добавляет прогон текста в конец абзаца
            if countOfTextEl < len(PhraseToMTK2):
                if PhraseToMTK2[countOfTextEl] == '1':
                    run.font.color.rgb = RGBColor(1, 0, 0)
                else:
                    run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                run.font.color.rgb = RGBColor(0, 0, 0)
            countOfTextEl += 1

    document.save('var077_modify.docx')

main()